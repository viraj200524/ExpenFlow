import google.generativeai as genai
import json
from docx import Document
from tabulate import tabulate
import yagmail

EXPENSE_POLICIES = {
    "Executive Level": {
        "Travel Expenses": {
            "Business Trips": "₹4,00,000 - ₹16,00,000 per month",
            "Local Transportation": "₹30,000 - ₹60,000 per month",
            "Mileage Reimbursement": "₹45,000 per month",
            "Parking Fees & Tolls": "Fully covered"
        },
        "Accommodation": {
            "Hotel Stays": "₹25,000 - ₹1,25,000 per night",
            "Rental Allowance": "₹2,50,000 - ₹6,00,000 per month",
            "Meals During Travel": "₹8,000 - ₹40,000 per day"
        },
        "Office Supplies and Equipment": {
            "Work Tools": "Unlimited (as per requirement)",
            "Home Office Setup": "Up to ₹8,00,000"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        },
        "Meals and Entertainment": {
            "Client Meetings": "Up to ₹4,00,000 per month",
            "Team Outings": "Up to ₹1,60,000 per month",
            "Daily Meal Allowance": "₹8,000 - ₹24,000 per day"
        }
    },
    "Senior Management": {
        "Travel Expenses": {
            "Business Trips": "₹1,50,000 - ₹4,00,000 per month",
            "Local Transportation": "₹20,000 - ₹25,000 per month",
            "Mileage Reimbursement": "₹30,000 per month",
            "Parking Fees & Tolls": "Up to ₹10,000 per month"
        },
        "Accommodation": {
            "Hotel Stays": "₹16,000 - ₹80,000 per night",
            "Rental Allowance": "₹1,00,000 - ₹1,50,000 per month",
            "Meals During Travel": "₹6,000 - ₹24,000 per day"
        },
        "Office Supplies and Equipment": {
            "Work Tools": "Up to ₹4,00,000 per year",
            "Home Office Setup": "Up to ₹4,00,000"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        },
        "Meals and Entertainment": {
            "Client Meetings": "Up to ₹2,40,000 per month",
            "Team Outings": "Up to ₹1,20,000 per month"
        }
    },
    "Middle Management": {
        "Travel Expenses": {
            "Business Trips": "₹1,60,000 - ₹5,60,000 per month",
            "Local Transportation": "₹16,000 - ₹20,000 per month",
            "Mileage Reimbursement": "₹25,000 per month"
        },
        "Accommodation": {
            "Rental Allowance": "₹80,000 - ₹90,000 per month"
        },
        "Office Supplies and Equipment": {
            "Work Tools": "Up to ₹2,40,000 per year",
            "Home Office Setup": "Up to ₹2,40,000"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        },
        "Meals and Entertainment": {
            "Client Meetings": "Up to ₹1,60,000 per event"
        }
    },
    "Lower Management": {
        "Travel Expenses": {
            "Business Trips": "₹80,000 - ₹4,00,000 per trip",
            "Local Transportation": "₹12,000 - ₹15,000 per month",
            "Mileage Reimbursement": "₹20,000 per month"
        },
        "Accommodation": {
            "Rental Allowance": "₹40,000 - ₹50,000 per month"
        },
        "Office Supplies and Equipment": {
            "Work Tools": "Up to ₹1,60,000 per year"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        }
    },
    "Team Leads & Supervisors": {
        "Travel Expenses": {
            "Business Trips": "₹40,000 - ₹2,40,000 per trip",
            "Local Transportation": "₹8,000 - ₹10,000 per month",
            "Mileage Reimbursement": "₹15,000 per month"
        },
        "Accommodation": {
            "Rental Allowance": "₹30,000 - ₹40,000 per month"
        },
        "Meals and Entertainment": {
            "Client Meetings": "Up to ₹80,000 per event"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        }
    },
    "Staff & Employees": {
        "Travel Expenses": {
            "Business Trips": "₹24,000 - ₹1,60,000 per trip",
            "Local Transportation": "Up to ₹5,000 per month",
            "Mileage Reimbursement": "₹10,000 per month"
        },
        "Accommodation": {
            "Rental Allowance": "₹15,000 - ₹30,000 per month"
        },
        "Office Supplies and Equipment": {
            "Work Tools": "Up to ₹80,000 per year"
        },
        "Communication Expenses": {
            "Mobile/Internet Bills": "Fully covered"
        }
    }
}

def generate_expense_report(json_data, api_key):
    """
    Generate professional expense reports (for employees and HR) using Gemini API,
    export them as .docx files, and send them via email using yagmail.
    Args:
        json_data (dict or list): The JSON data for receipts.
        api_key (str): API key for the Gemini model.
    Returns:
        None: Prints both reports directly and sends them via email.
    """
    # Set up Gemini API
    genai.configure(api_key=api_key)

    # Define expense categories
    expense_categories = [
        "Flight", "Hotel", "Stationery", "Travel", "Accommodation",
        "Office Supplies and Equipment", "Training and Development",
        "Health and Wellness", "Miscellaneous"
    ]

    # Helper function to categorize items
    def categorize_item(item_name):
        item_name_lower = item_name.lower()
        if "flight" in item_name_lower:
            return "Flight"
        elif "hotel" in item_name_lower:
            return "Hotel"
        elif "pen" in item_name_lower or "notebook" in item_name_lower:
            return "Stationery"
        elif "travel" in item_name_lower:
            return "Travel"
        elif "laptop" in item_name_lower or "printer" in item_name_lower:
            return "Office Supplies and Equipment"
        elif "training" in item_name_lower or "course" in item_name_lower:
            return "Training and Development"
        elif "gym" in item_name_lower or "medical" in item_name_lower:
            return "Health and Wellness"
        else:
            return "Miscellaneous"

    # Initialize report structure
    total_reimbursement = {cat: 0 for cat in expense_categories}
    total_non_reimbursable = {cat: 0 for cat in expense_categories}
    employee_breakdown = {}
    employee_non_reimbursable = {}
    violations_summary = {}

    # Validate json_data structure
    if isinstance(json_data, dict):
        json_data = [json_data]  # Convert single dictionary to list
    elif not isinstance(json_data, list):
        raise ValueError("json_data must be a dictionary or a list of dictionaries.")

    # Process data
    for receipt in json_data:
        if not isinstance(receipt, dict):
            print(f"Skipping invalid receipt (not a dictionary): {receipt}")
            continue

        # Extract relevant fields with error handling
        vendor_name = receipt.get("vendor", {}).get("name", "N/A") if isinstance(receipt.get("vendor"), dict) else "N/A"
        invoice_number = receipt.get("invoice_number", "N/A")
        total_amount = receipt.get("total", 0) or 0  # Ensure it's not None
        line_items = receipt.get("line_items", [])
        fraud_flags = receipt.get("meta", {}).get("fraud_flags", []) if isinstance(receipt.get("meta"), dict) else []
        employee_id = receipt.get("reference_number", "Unknown")

        # Initialize employee breakdowns if not already present
        if employee_id not in employee_breakdown:
            employee_breakdown[employee_id] = {cat: 0 for cat in expense_categories}
        if employee_id not in employee_non_reimbursable:
            employee_non_reimbursable[employee_id] = {
                cat: {"amount": 0, "violations": []} for cat in expense_categories
            }
        if employee_id not in violations_summary:
            violations_summary[employee_id] = []

        # Process line items
        for item in line_items:
            if not isinstance(item, dict):
                print(f"Skipping invalid line item (not a dictionary): {item}")
                continue

            item_name = item.get("description", "Unknown Item")
            item_total = item.get("total", 0) or 0  # Ensure it's not None
            category = categorize_item(item_name)

            # Use Gemini API to detect fraud flags
            def detect_fraud(item_description, total_amount, vendor_name):
                model = genai.GenerativeModel('gemini-pro')
                prompt = f"""
                Analyze the following expense for potential fraud or policy violations:
                Item Description: {item_description}
                Total Amount: {total_amount}
                Vendor Name: {vendor_name}
                Return a JSON object without bold with the following keys:
                "is_fraud": true/false (whether this expense is potentially fraudulent)
                "reason": Explanation of why it might be fraudulent (if applicable)
                """
                response = model.generate_content(prompt)
                try:
                    return json.loads(response.text)
                except json.JSONDecodeError:
                    return {"is_fraud": False, "reason": "Unable to determine fraud status."}

            fraud_analysis = detect_fraud(item_name, item_total, vendor_name)
            if fraud_analysis["is_fraud"]:
                fraud_flags.append(fraud_analysis["reason"])

            if fraud_flags:
                total_non_reimbursable[category] += item_total
                employee_non_reimbursable[employee_id][category]["amount"] += item_total
                employee_non_reimbursable[employee_id][category]["violations"].extend(fraud_flags)
                violations_summary[employee_id].extend(fraud_flags)
            else:
                total_reimbursement[category] += item_total
                employee_breakdown[employee_id][category] += item_total

    # Remove duplicate violations
    for employee_id in violations_summary:
        violations_summary[employee_id] = list(set(violations_summary[employee_id]))

    # Filter out zero-value categories
    def filter_zero_values(data_dict):
        return {k: v for k, v in data_dict.items() if v > 0}

    total_reimbursement = filter_zero_values(total_reimbursement)
    total_non_reimbursable = filter_zero_values(total_non_reimbursable)
    for employee_id in employee_breakdown:
        employee_breakdown[employee_id] = filter_zero_values(employee_breakdown[employee_id])
    for employee_id in employee_non_reimbursable:
        employee_non_reimbursable[employee_id] = {
            k: v for k, v in employee_non_reimbursable[employee_id].items() if v["amount"] > 0
        }

    # Prepare report data for summarization
    report_data = {
        "total_reimbursement": total_reimbursement,
        "total_non_reimbursable": total_non_reimbursable,
        "employee_breakdown": employee_breakdown,
        "employee_non_reimbursable": employee_non_reimbursable,
        "violations_summary": violations_summary
    }

    # Generate natural language summaries using Gemini API
    def generate_employee_report(report_data, employee_id):
        model = genai.GenerativeModel('gemini-pro')
        reimbursable = report_data['employee_breakdown'].get(employee_id, {})
        non_reimbursable = report_data['employee_non_reimbursable'].get(employee_id, {})
        violations = report_data['violations_summary'].get(employee_id, [])

        # Create a table for reimbursements
        reimbursement_table = [["Category", "Amount"]]
        for category, amount in reimbursable.items():
            reimbursement_table.append([category, f"₹{amount}"])

        # Create a table for non-reimbursable amounts
        non_reimbursable_table = [["Category", "Amount", "Violations"]]
        for category, details in non_reimbursable.items():
            non_reimbursable_table.append([category, f"₹{details['amount']}", ", ".join(details['violations'])])

        # Create a table for violations with policy details
        violations_table = [["Violation", "Policy"]]
        for violation in violations:
            policy = "Policy not found"
            for level, policies in EXPENSE_POLICIES.items():
                for category, limits in policies.items():
                    for key, value in limits.items():
                        if key.lower() in violation.lower():
                            policy = f"{key}: {value}"
                            break
            violations_table.append([violation, policy])

        # Generate textual summary
        prompt = f"""
        Generate a concise and professional expense report for an employee with ID {employee_id}.
        Include the following details:
        - Reimbursable Amounts by Category (in table format):
          {tabulate(reimbursement_table, headers="firstrow", tablefmt="grid")}
        - Non-Reimbursable Amounts by Category (in table format):
          {tabulate(non_reimbursable_table, headers="firstrow", tablefmt="grid")}
        - Violations Detected (in table format):
          {tabulate(violations_table, headers="firstrow", tablefmt="grid")}
        Provide personalized feedback and suggestions to help the employee avoid similar issues in the future.
        Use a friendly and motivational tone.
        """
        response = model.generate_content(prompt)
        return response.text

    def generate_hr_report(report_data):
        model = genai.GenerativeModel('gemini-pro')

        # Create a table for total reimbursements
        total_reimbursement_table = [["Category", "Amount"]]
        for category, amount in report_data['total_reimbursement'].items():
            total_reimbursement_table.append([category, f"₹{amount}"])

        # Create a table for total non-reimbursable amounts
        total_non_reimbursable_table = [["Category", "Amount"]]
        for category, amount in report_data['total_non_reimbursable'].items():
            total_non_reimbursable_table.append([category, f"₹{amount}"])

        # Create a table for employee-wise reimbursements
        employee_reimbursement_table = [["Employee ID", "Category", "Amount"]]
        for employee_id, categories in report_data['employee_breakdown'].items():
            for category, amount in categories.items():
                employee_reimbursement_table.append([employee_id, category, f"₹{amount}"])

        # Create a table for employee-wise non-reimbursable amounts
        employee_non_reimbursable_table = [["Employee ID", "Category", "Amount", "Violations"]]
        for employee_id, categories in report_data['employee_non_reimbursable'].items():
            for category, details in categories.items():
                employee_non_reimbursable_table.append(
                    [employee_id, category, f"₹{details['amount']}", ", ".join(details['violations'])]
                )

        # Create a table for violations with policy details
        violations_table = [["Employee ID", "Violation", "Policy"]]
        for employee_id, violations in report_data['violations_summary'].items():
            for violation in violations:
                policy = "Policy not found"
                for level, policies in EXPENSE_POLICIES.items():
                    for category, limits in policies.items():
                        for key, value in limits.items():
                            if key.lower() in violation.lower():
                                policy = f"{key}: {value}"
                                break
                violations_table.append([employee_id, violation, policy])

        # Generate textual summary
        prompt = f"""
        Generate a detailed and professional expense report for HR.
        Include the following details:
        - Total Reimbursement by Category (in table format):
          {tabulate(total_reimbursement_table, headers="firstrow", tablefmt="grid")}
        - Total Non-Reimbursable Amounts by Category (in table format):
          {tabulate(total_non_reimbursable_table, headers="firstrow", tablefmt="grid")}
        - Employee-wise Breakdown of Reimbursable Amounts (in table format):
          {tabulate(employee_reimbursement_table, headers="firstrow", tablefmt="grid")}
        - Employee-wise Breakdown of Non-Reimbursable Amounts (in table format):
          {tabulate(employee_non_reimbursable_table, headers="firstrow", tablefmt="grid")}
        - Employee-wise Violations (in table format):
          {tabulate(violations_table, headers="firstrow", tablefmt="grid")}
        Highlight compliance issues, flagged items, and provide actionable recommendations for improving expense management.
        Use a formal and detailed tone.
        """
        response = model.generate_content(prompt)
        return response.text

    # Generate reports
    employee_reports = {}
    for employee_id in employee_breakdown.keys():
        employee_reports[employee_id] = generate_employee_report(report_data, employee_id)

    hr_report_text = generate_hr_report(report_data)

    # Export reports to .docx format
    def export_to_docx(report_text, filename):
        doc = Document()
        doc.add_heading("Expense Report", level=1)
        doc.add_paragraph(report_text)
        doc.save(filename)

    for employee_id, report_text in employee_reports.items():
        export_to_docx(report_text, f"employee_report_{employee_id}.docx")

    export_to_docx(hr_report_text, "hr_report.docx")

    # Send emails with attachments using yagmail
    def send_email(sender_email, recipient_email, subject, body, attachment_path):
        # Initialize yagmail SMTP connection
        yag = yagmail.SMTP(sender_email, "btjr mnzc ozto ntcg")  # Replace with your App Password
        # Send email with attachment
        yag.send(
            to=recipient_email,
            subject=subject,
            contents=body,
            attachments=attachment_path
        )

    # Send HR report
    send_email(
        sender_email="virajv2005@gmail.com",
        recipient_email="vrvora_b23@ce.vjti.ac.in",
        subject="MONTHLY COMPANY EXPENSE REPORT FOR HR",
        body="Please find attached the monthly expense report for HR.",
        attachment_path="hr_report.docx"
    )

    # Send Employee reports
    for employee_id, report_text in employee_reports.items():
        send_email(
            sender_email="virajv2005@gmail.com",
            recipient_email="vrvora_b23@ce.vjti.ac.in",  # Replace with actual employee email logic
            subject="MONTHLY COMPANY EXPENSE REPORT FOR EMPLOYEE",
            body="Please find attached the monthly expense report for your review.",
            attachment_path=f"employee_report_{employee_id}.docx"
        )

    print("\n--- Reports Generated and Sent Successfully ---\n")